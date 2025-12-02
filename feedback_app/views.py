from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.decorators import login_required
from django.contrib.auth.models import User
from django.db.models import Avg, Count
from .models import Faculty, Student, Subject, Section, StudentSubject, StudentFacultyMap, Feedback

# Login View
from django.contrib.auth import authenticate, login
from django.contrib import messages
from django.shortcuts import render, redirect
from .models import Student, Faculty

def login_view(request):
    if request.method == 'POST':
        username = request.POST.get('username', '').strip()
        password = request.POST.get('password', '').strip()

        user = authenticate(request, username=username, password=password)

        if user:
            login(request, user)

            if user.is_superuser:
                return redirect('admin_dashboard')

            elif Faculty.objects.filter(user=user.username).exists():
                return redirect('faculty_dashboard')

            elif Student.objects.filter(user=user.username).exists():
                return redirect('student_dashboard')

            else:
                messages.error(request, "User exists but not assigned a role.")
                return redirect('login')

        else:
            messages.error(request, "Invalid username or password.")
            return render(request, 'feedback_app/login.html')

    return render(request, 'feedback_app/login.html')


# Logout View
@login_required
def logout_view(request):
    logout(request)
    return redirect('login')

# Admin Dashboard
@login_required
def admin_dashboard(request):
    selected_year = request.GET.get('year')
    
    all_students = Student.objects.all()
    years = sorted(set(student.year for student in all_students))

    if selected_year:
        students = Student.objects.filter(year=selected_year).order_by('usn')  # 🟢 Sorted filtered students
    else:
        students = Student.objects.all().order_by('usn')  # 🟢 Default to all students
    
    faculty = Faculty.objects.all().order_by('name')
    subjects = Subject.objects.all()
    sections = Section.objects.all()

    return render(request, 'feedback_app/admin_dashboard.html', {
        'students': students,
        'years': years,
        'selected_year': selected_year,
        'faculty': faculty,
        'subjects': subjects,
        'sections': sections
    })




from django.shortcuts import render, redirect
from django.http import HttpResponse
from django.contrib import messages
from django.db.models import Count
from .models import Faculty, Subject, Feedback, FeedbackQuestion, StudentSubject, Student, FacultySubject
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
from datetime import datetime
import os
from django.conf import settings

@login_required
def faculty_dashboard(request):
    username = request.user.username
    faculty = Faculty.objects.filter(user=username).first()
    if not faculty:
        messages.error(request, "Faculty profile not found.")
        return redirect("login")

    # Subjects assigned to this faculty (by name)
    assigned_subject_names = FacultySubject.objects.filter(
        faculty_name=faculty.name
    ).values_list('subject_name', flat=True).distinct()
    selected_subjects = Subject.objects.filter(name__in=assigned_subject_names)

    semesters = ['1', '2', '3', '4']
    sections = ['All', 'A', 'B', 'C']  # keep "All"
    years = ['2019', '2020', '2021', '2022', '2023', '2024', '2025']

    # GET filters
    semester = request.GET.get('semester')
    section = request.GET.get('section')   # may be 'All' or specific
    year = request.GET.get('year')
    selected_subject_id = request.GET.get('subject_id')
    download = request.GET.get('download')

    # Show subjects only if filters are all selected (you used this earlier)
    filtered_subjects = Subject.objects.none()
    if semester and section and year:
        filtered_subjects = selected_subjects.filter(semester=semester)

    selected_subject = Subject.objects.filter(id=selected_subject_id).first() if selected_subject_id else None

    # Initialize outputs
    feedback_table = []
    total_students = 0
    present_students = 0
    overall_total = 0
    overall_avg = 0
    interpretation = ""
    feedback_list = Feedback.objects.none()

    if selected_subject:
        # -- Determine effective FacultySubject rows for this faculty+subject --
        fs_qs = FacultySubject.objects.filter(
            faculty_name=faculty.name,
            subject_name=selected_subject.name
        )
        # Narrow by semester/year if filters provided (so we consider the assignment in that term)
        if semester:
            fs_qs = fs_qs.filter(semester=semester)
        if year:
            fs_qs = fs_qs.filter(year=year)

        # If user explicitly chose a specific section (not "All"), respect that
        # Otherwise we derive sections from faculty assignments
        if section and section != "All":
            effective_sections = [section]
        else:
            effective_sections = list(fs_qs.values_list('section', flat=True).distinct()) if fs_qs.exists() else []

        # If no FacultySubject rows found, fallback to faculty.section(s) inference
        if not effective_sections:
            if hasattr(faculty, 'section') and faculty.section:
                effective_sections = [faculty.section]
            elif hasattr(faculty, 'sections'):
                effective_sections = list(faculty.sections.all().values_list('name', flat=True))
            else:
                effective_sections = []

        # Build students_qs according to effective_sections
        students_qs = Student.objects.none()

        if 'All' in effective_sections:
            # All sections: every student who is assigned this subject (across sections)
            assigned_student_names = StudentSubject.objects.filter(
                subject_name=selected_subject.name
            ).values_list('student_name', flat=True).distinct()

            students_qs = Student.objects.filter(name__in=assigned_student_names).distinct()

            # apply semester/year filters if provided (Student model fields)
            if semester:
                students_qs = students_qs.filter(semester=semester)
            if year:
                students_qs = students_qs.filter(year=year)
        else:
            # Only students in the effective_sections and assigned to this subject
            if effective_sections:
                names_in_sections = Student.objects.filter(
                    section__in=effective_sections
                ).values_list('name', flat=True).distinct()

                assigned_student_names = StudentSubject.objects.filter(
                    subject_name=selected_subject.name,
                    student_name__in=names_in_sections
                ).values_list('student_name', flat=True).distinct()

                students_qs = Student.objects.filter(name__in=assigned_student_names).distinct()

                # apply semester/year filters if provided
                if semester:
                    students_qs = students_qs.filter(semester=semester)
                if year:
                    students_qs = students_qs.filter(year=year)
            else:
                students_qs = Student.objects.none()

        total_students = students_qs.count()

        # Feedback list restricted to assigned students (keeps present_students consistent)
        feedback_list = Feedback.objects.filter(
            faculty=faculty,
            subject=selected_subject,
            student__in=students_qs
        )

        present_students = feedback_list.values('student').distinct().count()

        # Questions breakdown (same as feedback_report)
        questions = FeedbackQuestion.objects.all()
        overall_total = 0

        for idx, question in enumerate(questions, start=1):
            feedbacks = feedback_list.filter(question=question)

            # If user selected specific section (and not All) filter feedbacks accordingly
            if section and section != "All":
                feedbacks = feedbacks.filter(student__section=section)
            elif section == "All":
                # When "All" is selected, feedbacks already limited by students_qs (which is across sections)
                pass

            # apply year filter on feedback student records as well
            if year:
                feedbacks = feedbacks.filter(student__year=year)

            counts = {i: feedbacks.filter(rating=i).count() for i in range(1, 6)}
            total_responses = sum(counts.values())

            weighted_total = (
                counts[1] * 20 +
                counts[2] * 40 +
                counts[3] * 60 +
                counts[4] * 80 +
                counts[5] * 100
            )

            avg_score = round(weighted_total / total_responses, 2) if total_responses else 0
            overall_total += avg_score

            feedback_table.append({
                'sl_no': idx,
                'question': question.text,
                'counts': counts,
                'total': weighted_total,
                'average': avg_score,
                'marks': avg_score,
            })

        overall_avg = round(overall_total / len(questions), 2) if questions else 0

        if overall_avg >= 85:
            interpretation = "Excellent"
        elif overall_avg >= 75:
            interpretation = "Very Good"
        elif overall_avg >= 60:
            interpretation = "Good"
        elif overall_avg >= 50:
            interpretation = "Satisfactory"
        else:
            interpretation = "Poor"

        # Download Word report (same as your previous implementation)
        if download:
            doc = Document()
            logo_path = os.path.join(settings.BASE_DIR, 'feedback_app', 'static', 'images', 'logo.png')
            if os.path.exists(logo_path):
                doc.add_picture(logo_path, width=Inches(1.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

            run = doc.add_paragraph().add_run("AMC ENGINEERING COLLEGE")
            run.bold = True
            run.font.size = Pt(16)
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph("Bannerghatta Road, Bengaluru, Karnataka - 560083").alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph("DEPARTMENT OF COMPUTER APPLICATIONS").alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph("FACULTY FEEDBACK REPORT").alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph("Rating: 1–Poor  2–Satisfactory  3–Good  4–Very Good  5–Excellent").alignment = WD_ALIGN_PARAGRAPH.CENTER

            info_table = doc.add_table(rows=2, cols=3)
            info_table.style = 'Table Grid'
            info_table.autofit = False

            info_table.cell(0, 0).text = f"Faculty: {faculty.name}"
            info_table.cell(0, 1).text = f"Faculty ID: {faculty.usn if hasattr(faculty, 'usn') else ''}"
            info_table.cell(0, 2).text = f"Subject: {selected_subject.name}"
            info_table.cell(1, 0).text = f"Semester: {faculty.semester}"
            info_table.cell(1, 1).text = f"Strength of the class: {total_students}"
            info_table.cell(1, 2).text = f"Student Present: {present_students}"

            doc.add_paragraph()

            table = doc.add_table(rows=1, cols=9)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text = "Sl No"
            hdr[1].text = "Feedback Category"
            hdr[2].text = "1"
            hdr[3].text = "2"
            hdr[4].text = "3"
            hdr[5].text = "4"
            hdr[6].text = "5"
            hdr[7].text = "Total"
            hdr[8].text = "Marks"

            widths = [
                Inches(0.3),
                Inches(6.4),
                Inches(0.1),
                Inches(0.1),
                Inches(0.1),
                Inches(0.1),
                Inches(0.1),
                Inches(0.3),
                Inches(0.3),
            ]
            # set widths for header row (docx widths are hints; this keeps layout similar)
            for row in table.rows:
                for i, width in enumerate(widths):
                    try:
                        row.cells[i].width = width
                    except Exception:
                        pass

            for row in feedback_table:
                r = table.add_row().cells
                r[0].text = str(row['sl_no'])
                r[1].text = row['question']
                for i in range(1, 6):
                    r[i + 1].text = str(row['counts'][i])
                r[7].text = str(row['total'])
                r[8].text = str(row['marks'])

            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = para.add_run(f"Average Marks: {overall_avg}")
            run.bold = True
            run.font.size = Pt(16)

            doc.add_paragraph("\nTotal=(20%*no of 1 rating)+(40%*no of 2 rating)+(60%*no of 3 rating)+(80%*no of 4 rating)+(100%*no of 5 rating)")
            doc.add_paragraph("Marks=Total/No of Students Present")
            doc.add_paragraph("Poor:<50%, Satisfactory:>=50% & <65%, Good:>=60% & <75%, Very Good:>=75% & <85%, Excellent:>=85")
            doc.add_paragraph(f"Overall Feedback: {interpretation}")
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            paragraph.add_run("\n\nHOD, MCA").bold = True
            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            filename = f"{faculty.name}_Feedback_{selected_subject.code if hasattr(selected_subject, 'code') else selected_subject.name}.docx"
            return HttpResponse(
                buf.read(),
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                headers={'Content-Disposition': f'attachment; filename={filename}'}
            )

    return render(request, 'feedback_app/faculty_dashboard.html', {
        'faculty': faculty,
        'subjects': selected_subjects,
        'selected_subjects': selected_subjects,
        'filtered_subjects': filtered_subjects,
        'selected_subject': selected_subject,
        'feedback_list': feedback_list,
        'feedback_table': feedback_table,
        'total_students': total_students,
        'present_students': present_students,
        'overall_avg': overall_avg,
        'interpretation': interpretation,
        'semesters': semesters,
        'sections': sections,
        'years': years,
    })





# Student Dashboard
from django.shortcuts import render, redirect
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from .models import (
    Student, Faculty, Subject, FeedbackQuestion, Feedback,
    StudentSubject, Facility, FacilityFeedback,
    CourseEndQuestion, CourseEndFeedback
)
@login_required
def student_dashboard(request):
    username = request.user.username
    student = Student.objects.filter(user=username).first()
    if not student:
        return redirect('login')

    # Get all subjects mapped to the student
    subject_names = StudentSubject.objects.filter(student_name=student.name).values_list('subject_name', flat=True)
    subjects = Subject.objects.filter(name__in=subject_names)

    selected_subject = None
    selected_faculty = None
    selected_faculties = []
    questions = []
    show_form = False
    final_comments = ''
    ratings = ['1', '2', '3', '4', '5']

    if request.method == 'POST':
        selected_subject = request.POST.get('subject')
        selected_faculty = request.POST.get('faculty')
        final_comments = request.POST.get('final_comments', '')

        if selected_subject:
            faculty_names = FacultySubject.objects.filter(subject_name=selected_subject).values_list('faculty_name', flat=True)
            selected_faculties = Faculty.objects.filter(name__in=faculty_names)

        # If ratings are included, assume form is submitted
        if selected_subject and selected_faculty and any(f"rating_{q.id}" in request.POST for q in FeedbackQuestion.objects.all()):
            questions = FeedbackQuestion.objects.all()
            submitted_ratings = []
            faculty_obj = Faculty.objects.filter(name=selected_faculty).first()
            subject_obj = Subject.objects.filter(name=selected_subject).first()

            for question in questions:
                rating = request.POST.get(f"rating_{question.id}")
                if rating:
                    rating_int = int(rating)
                    submitted_ratings.append(rating_int)
                    Feedback.objects.create(
                        student=student,
                        faculty=faculty_obj,
                        subject=subject_obj,
                        question=question,
                        rating=rating_int,
                        comments=final_comments
                    )

            # Save feedback to Excel
            save_feedback_to_excel(request, student, questions, submitted_ratings)

            return redirect('course_end_feedback')

        elif selected_subject and selected_faculty:
            questions = FeedbackQuestion.objects.all()
            show_form = True

    elif request.method == 'GET':
        selected_subject = request.GET.get('subject')
        if selected_subject:
            faculty_names = FacultySubject.objects.filter(subject_name=selected_subject).values_list('faculty_name', flat=True)
            selected_faculties = Faculty.objects.filter(name__in=faculty_names)

    return render(request, 'feedback_app/student_dashboard.html', {
        'student': student,
        'subjects': subjects,
        'selected_subject': selected_subject,
        'selected_faculty': selected_faculty,
        'selected_faculties': selected_faculties,
        'questions': questions,
        'ratings': ratings,
        'show_form': show_form,
    })



    
import os
from openpyxl import Workbook, load_workbook
from openpyxl.utils import get_column_letter
from datetime import datetime

def save_feedback_to_excel(request, student, questions, ratings, sheet_name='Feedback'):
    import os
    from openpyxl import Workbook, load_workbook
    from django.contrib import messages
    from django.conf import settings

    year = str(student.year) if student.year else "UnknownYear"
    folder_path = os.path.join(settings.BASE_DIR, 'feedback_excel')
    os.makedirs(folder_path, exist_ok=True)
    file_path = os.path.join(folder_path, f"{year}.xlsx")

    try:
        # Create workbook if it doesn't exist
        if not os.path.exists(file_path):
            wb = Workbook()
            ws = wb.active
            ws.title = sheet_name
            headers = ['UserID'] + [f"Q{q.id}" for q in questions] + ['Total', 'Average']
            ws.append(headers)

            # Append question details only once
            ws.append([])
            ws.append(["Q.ID", "Question Text"])
            for q in questions:
                ws.append([f"Q{q.id}", q.text])

            wb.save(file_path)

        # Load workbook
        wb = load_workbook(file_path)

        # Create or select sheet
        if sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
        else:
            ws = wb.create_sheet(title=sheet_name)
            headers = ['UserID'] + [f"Q{q.id}" for q in questions] + ['Total', 'Average']
            ws.append(headers)
            ws.append([])
            ws.append(["Q.ID", "Question Text"])
            for q in questions:
                ws.append([f"Q{q.id}", q.text])

        # Find the row index where actual feedback ends
        # Start appending before the empty row before "Q.ID"
        insert_row = ws.max_row + 1
        for i in range(1, ws.max_row + 1):
            if ws.cell(i, 1).value == "Q.ID":
                insert_row = i - 2
                break

        # Shift question text section down by 1 row
        ws.insert_rows(insert_row + 1)

        # Write feedback entry
        total = sum(ratings)
        avg = round(total / len(ratings), 2) if ratings else 0
        ws.cell(row=insert_row + 1, column=1, value=student.user)
        for i, rating in enumerate(ratings):
            ws.cell(row=insert_row + 1, column=2 + i, value=rating)
        ws.cell(row=insert_row + 1, column=2 + len(ratings), value=total)
        ws.cell(row=insert_row + 1, column=3 + len(ratings), value=avg)

        # Only apply weighted average for 'Course End Feedback'
        if sheet_name == 'Course End Feedback':
            all_ratings = []
            for row in ws.iter_rows(min_row=2, max_row=insert_row, min_col=2, max_col=1 + len(questions)):
                for cell in row:
                    if isinstance(cell.value, int):
                        all_ratings.append(cell.value)

            count_1 = all_ratings.count(1)
            count_2 = all_ratings.count(2)
            count_3 = all_ratings.count(3)
            count_4 = all_ratings.count(4)
            count_5 = all_ratings.count(5)

            weighted_total = (
                count_1 * 20 +
                count_2 * 40 +
                count_3 * 60 +
                count_4 * 80 +
                count_5 * 100
            )
            num_students = insert_row - 1  # Exclude header
            weighted_avg = round(weighted_total / num_students, 2) if num_students else 0

            ws.insert_rows(insert_row + 2)
            ws.cell(row=insert_row + 2, column=1, value="Weighted Total")
            ws.cell(row=insert_row + 2, column=2 + len(ratings), value=weighted_total)
            ws.cell(row=insert_row + 2, column=3 + len(ratings), value=weighted_avg)

        wb.save(file_path)

    except PermissionError:
        messages.error(request, f"Permission denied: Cannot save feedback to {file_path}. Please close the Excel file if it is open.")


@login_required
def course_end_feedback(request):
    student = Student.objects.filter(user=request.user.username).first()
    if not student:
        return redirect('login')

    questions = CourseEndQuestion.objects.all()

    if request.method == 'POST':
        ratings = []
        for question in questions:
            rating = request.POST.get(f"rating_{question.id}")
            if rating:
                rating_int = int(rating)
                ratings.append(rating_int)
                CourseEndFeedback.objects.create(
                    student=student,
                    question=question,
                    rating=rating_int
                )

        # Save to Excel
        save_feedback_to_excel(request, student, questions, ratings, sheet_name='CourseEndFeedback')

        return redirect('facility_feedback')  # move to next feedback section

    return render(request, 'feedback_app/course_end_feedback.html', {
        'student': student,
        'questions': questions,
        'ratings': [1, 2, 3, 4, 5],
    })

@login_required
def facility_feedback(request):
    student = Student.objects.filter(user=request.user.username).first()
    if not student:
        return redirect('login')

    facilities = Facility.objects.all()

    if request.method == 'POST':
        ratings = []
        for facility in facilities:
            rating = request.POST.get(f"rating_{facility.id}")
            if rating:
                FacilityFeedback.objects.create(
                    student=student,
                    facility=facility,
                    rating=int(rating)
                )
                ratings.append(int(rating))

        # Save to Excel in FacilityFeedback sheet
        save_feedback_to_excel(request, student, facilities, ratings, sheet_name='FacilityFeedback')

        return redirect('admin_dashboard')

    return render(request, 'feedback_app/facility_feedback.html', {
        'student': student,
        'facilities': facilities,
        'ratings': [1, 2, 3, 4, 5],
    })




from django.db import IntegrityError
from django.contrib import messages
# Add Faculty
@login_required
def add_faculty(request):
    if request.method == "POST":
        username = request.POST.get("username")
        password = request.POST.get("password")
        usn = request.POST.get("usn")

        if not username or not password or not usn:
            messages.error(request, "All fields are required!")
            return redirect("add_faculty")

        # Check if username already exists in User
        if User.objects.filter(username=username).exists():
            messages.error(request, "Username already exists!")
            return redirect("add_faculty")

        # Check if USN already exists in Faculty
        if Faculty.objects.filter(usn=usn).exists():
            messages.error(request, "Faculty ID already exists!")
            return redirect("add_faculty")

        try:
            # Create Django User
            user = User.objects.create_user(username=username, password=password)

            # Create Faculty object
            Faculty.objects.create(
                user=username,   # store username as string
                name=username,   # or use a separate 'name' field from form
                usn=usn
            )

            messages.success(request, f"Faculty '{username}' added successfully!")
            return redirect("admin_dashboard")

        except IntegrityError as e:
            messages.error(request, f"Error creating faculty: {str(e)}")
            return redirect("add_faculty")

    return render(request, "feedback_app/add_faculty.html")

from django.shortcuts import render, redirect
from django.contrib import messages
from .models import Student
from openpyxl import load_workbook
import random
import string
from django.contrib.auth.models import User  # ✅ Import User
from django.contrib.auth.decorators import login_required

def generate_random_password(length=8):
    chars = string.ascii_letters + string.digits
    return ''.join(random.choices(chars, k=length))

@login_required
def add_student(request):
    if request.method == 'POST':
        form_type = request.POST.get('form_type')

        # ================= MANUAL ENTRY =================
        if form_type == 'manual':
            name = request.POST.get('name').strip()
            usn = request.POST.get('usn').strip()
            section = request.POST.get('section', 'A')
            year = int(request.POST.get('year', 2020))
            semester = request.POST.get('semester', '1')

            if not name or not usn:
                messages.error(request, "Name and USN are required.")
                return redirect('add_student')

            student, created = Student.objects.get_or_create(usn=usn)
            student.name = name
            student.section = section
            student.year = year
            student.semester = semester

            if created or not student.generated_userid:
                # Generate unique MCA ID
                existing_ids = set(Student.objects.values_list('generated_userid', flat=True))
                while True:
                    new_id = f"MCA{random.randint(1, 999):03}"
                    if new_id not in existing_ids:
                        break

                password = generate_random_password()
                student.generated_userid = new_id
                student.generated_password = password

                # ✅ Create Django User
                if User.objects.filter(username=new_id).exists():
                    User.objects.filter(username=new_id).delete()  # Avoid duplicate
                User.objects.create_user(username=new_id, password=password, first_name=name)

                student.user = new_id  # Link back
            student.save()
            messages.success(request, f"Student added with ID {student.generated_userid}")
            return redirect('add_student')

        # ================= EXCEL UPLOAD =================
        elif form_type == 'excel' and 'file' in request.FILES:
            excel_file = request.FILES['file']
            wb = load_workbook(filename=excel_file)
            ws = wb.active

            new_students = []
            updated_count = 0

            for row in ws.iter_rows(min_row=2, values_only=True):
                if not row or len(row) < 2 or not row[0] or not row[1]:
                    continue

                usn = str(row[0]).strip()
                name = str(row[1]).strip()
                section = str(row[2]).strip() if len(row) > 2 and row[2] else "A"
                year = int(row[3]) if len(row) > 3 and row[3] else 2020
                semester = str(row[4]) if len(row) > 4 and row[4] else "1"

                student = Student.objects.filter(usn=usn).first()

                if student:
                    student.name = name
                    student.section = section
                    student.year = year
                    student.semester = semester
                    student.save()
                    updated_count += 1
                else:
                    new_students.append({
                        "usn": usn,
                        "name": name,
                        "section": section,
                        "year": year,
                        "semester": semester
                    })

            # Generate MCA IDs and create users
            for i, data in enumerate(new_students):
                password = generate_random_password()
                while True:
                    mca_id = f"MCA{random.randint(1, 999):03}"
                    if not User.objects.filter(username=mca_id).exists():
                        break

                # Create Django User
                User.objects.create_user(username=mca_id, password=password, first_name=data['name'])

                # Create Student record
                Student.objects.create(
                    user=mca_id,
                    usn=data['usn'],
                    name=data['name'],
                    section=data['section'],
                    year=data['year'],
                    semester=data['semester'],
                    generated_userid=mca_id,
                    generated_password=password
                )

            messages.success(request, f"{len(new_students)} students added, {updated_count} updated.")
            return redirect('add_student')

    return render(request, 'feedback_app/add_student.html')


# Add Subject
@login_required
def add_subject(request):
    if request.method == 'POST':
        code = request.POST.get('code')
        name = request.POST.get('name')
        Subject.objects.create(code=code, name=name)
        return redirect('admin_dashboard')

    return render(request, 'feedback_app/add_subject.html')

# Assign Subject
from .models import Student, Subject, Faculty, StudentSubject, FacultySubject
@login_required
def assign_subject(request):
    if request.method == 'POST':
        assign_type = request.POST.get('assign_type')

        if assign_type == 'student':
            subject_id = request.POST.get('subject_id_student')
            student_ids = request.POST.getlist('student_ids')
            subject = Subject.objects.get(id=subject_id)

            for sid in student_ids:
                student = Student.objects.get(id=sid)
                StudentSubject.objects.create(
                    student_name=student.name,
                    subject_name=subject.name
                )

        elif assign_type == 'faculty':
            subject_id = request.POST.get('subject_id_faculty')
            faculty_id = request.POST.get('faculty_id')
            semester = request.POST.get('semester')
            section = request.POST.get('section')
            year = request.POST.get('year')

            faculty = Faculty.objects.get(id=faculty_id)
            subject = Subject.objects.get(id=subject_id)

            if section == "All":
                # Get all sections dynamically
                all_sections = Student.objects.values_list('section', flat=True).distinct()
                duplicate_found = False

                for sec in all_sections:
                    if FacultySubject.objects.filter(
                        faculty_name=faculty.name,
                        subject_name=subject.name,
                        semester=semester,
                        section=sec,
                        year=year
                    ).exists():
                        duplicate_found = True
                    else:
                        FacultySubject.objects.create(
                            faculty_name=faculty.name,
                            subject_name=subject.name,
                            semester=semester,
                            section=sec,
                            year=year
                        )

                if duplicate_found:
                    messages.warning(request, "Some sections already had this subject assigned.")

            else:
                # Check for duplicate before creating
                if FacultySubject.objects.filter(
                    faculty_name=faculty.name,
                    subject_name=subject.name,
                    semester=semester,
                    section=section,
                    year=year
                ).exists():
                    messages.error(request, "Subject already assigned for this faculty and section.")
                else:
                    FacultySubject.objects.create(
                        faculty_name=faculty.name,
                        subject_name=subject.name,
                        semester=semester,
                        section=section,
                        year=year
                    )

        return redirect('admin_dashboard')

    students = Student.objects.all()
    faculties = Faculty.objects.all()
    subjects = Subject.objects.all()
    semesters = ['1', '2', '3', '4', '5', '6']
    sections = list(Student.objects.values_list('section', flat=True).distinct()) + ['All']
    years = ['2023', '2024', '2025']

    return render(request, 'feedback_app/assign_subject.html', {
        'students': students,
        'subjects': subjects,
        'faculties': faculties,
        'semesters': semesters,
        'sections': sections,
        'years': years,
    })

# Assign Section
@login_required
def assign_section(request):
    if request.method == 'POST':
        section_id = request.POST.get('section')
        semester = request.POST.get('semester')
        year = int(request.POST.get('year'))

        selected_section = Section.objects.get(id=section_id)

        student_ids = request.POST.getlist('student[]')
        faculty_ids = request.POST.getlist('faculty[]')

        for sid in student_ids:
            student = Student.objects.get(id=sid)
            student.section = selected_section.name
            student.year = year
            student.semester = semester  # 🔥 Add semester assignment
            student.save()

        for fid in faculty_ids:
            faculty = Faculty.objects.get(id=fid)
            faculty.section = selected_section.name
            faculty.year = year
            faculty.semester = semester  # 🔥 Add semester assignment
            faculty.save()

        return redirect('admin_dashboard')

    students = Student.objects.all()
    faculties = Faculty.objects.all()
    sections = Section.objects.all()


    return render(request, 'feedback_app/assign_section.html', {
        'students': students,
        'faculties': faculties,
        'sections': sections,
        'years': range(2019, 2026),
        'semesters': ['1', '2', '3', '4', '5', '6']
    })

import re
from django.db import transaction
from .models import Student, Faculty, StudentFacultyMap

# Map Student to Faculty
@login_required
def map_student_faculty(request):
    selected_year = request.GET.get('year')

    students = Student.objects.filter(year=selected_year).order_by('usn') if selected_year else Student.objects.all().order_by('usn')
    faculties = Faculty.objects.filter(year=selected_year).order_by('name') if selected_year else Faculty.objects.all().order_by('name')

    if request.method == 'POST':
        # delete old mappings for that year
        if selected_year:
            StudentFacultyMap.objects.filter(student__year=selected_year).delete()

        faculty_ids = request.POST.get('faculty_order', '').split(',')
        faculty_queryset = Faculty.objects.filter(id__in=faculty_ids)
        faculty_dict = {str(f.id): f for f in faculty_queryset}
        faculty_list_ordered = [faculty_dict[fac_id] for fac_id in faculty_ids if fac_id in faculty_dict]

        student_list = list(students)
        student_count = len(student_list)
        faculty_count = len(faculty_list_ordered)

        if faculty_count == 0:
            return HttpResponse("No faculty selected for mapping.")

        chunk_size = student_count // faculty_count
        remainder = student_count % faculty_count

        index = 0
        with transaction.atomic():
            for i, faculty in enumerate(faculty_list_ordered):
                count = chunk_size + (1 if i < remainder else 0)
                for _ in range(count):
                    if index < student_count:
                        StudentFacultyMap.objects.create(
                            student=student_list[index],
                            faculty=faculty
                        )
                        index += 1

        return redirect(f"{request.path}?year={selected_year}")

    mappings = StudentFacultyMap.objects.select_related('student', 'faculty')\
        .filter(student__year=selected_year).order_by('student__usn') if selected_year else []

    return render(request, 'feedback_app/map_student_faculty.html', {
        'students': students,
        'faculties': faculties,
        'mappings': mappings,
        'selected_year': selected_year,
        'year_choices': Student.objects.values_list('year', flat=True).distinct().order_by('year'),
    })




# Feedback Report
from django.db.models import Avg, Count

@login_required
def feedback_report(request):
    # Step 1: Summary table for all faculty/subjects
    report = Feedback.objects.values(
        'faculty__id', 'faculty__name',
        'subject__id', 'subject__name'
    ).annotate(
        avg_rating=Avg('rating') * 20,  # Convert 1–5 to percentage
        total=Count('id')
    ).order_by('faculty__name')

    # Step 2: Get selected faculty and subject from GET params
    faculty_id = request.GET.get('faculty_id')
    subject_id = request.GET.get('subject_id')

    feedback_table = []
    total_students = 0
    present_students = 0
    overall_avg = 0
    interpretation = ""
    selected_faculty = None
    selected_subject = None

    if faculty_id and subject_id:
        selected_faculty = Faculty.objects.filter(id=faculty_id).first()
        selected_subject = Subject.objects.filter(id=subject_id).first()

        if selected_faculty and selected_subject:
            # --- Determine assigned students based on FacultySubject rows ---
            faculty_subjects = FacultySubject.objects.filter(
                faculty_name=selected_faculty.name,
                subject_name=selected_subject.name
            )

            students_qs = Student.objects.none()

            if faculty_subjects.exists():
                sections = list(faculty_subjects.values_list('section', flat=True).distinct())

                if 'All' in sections:
                    # All sections: every student who is assigned this subject (across sections)
                    assigned_student_names = StudentSubject.objects.filter(
                        subject_name=selected_subject.name
                    ).values_list('student_name', flat=True).distinct()

                    students_qs = Student.objects.filter(name__in=assigned_student_names).distinct()
                else:
                    # Only the sections present in FacultySubject rows
                    # get student names in those sections
                    student_names_in_sections = Student.objects.filter(
                        section__in=sections
                    ).values_list('name', flat=True).distinct()

                    assigned_student_names = StudentSubject.objects.filter(
                        subject_name=selected_subject.name,
                        student_name__in=student_names_in_sections
                    ).values_list('student_name', flat=True).distinct()

                    students_qs = Student.objects.filter(name__in=assigned_student_names).distinct()

                total_students = students_qs.count()
            else:
                # Fallback: if no FacultySubject record, try to infer from faculty->section(s)
                if hasattr(selected_faculty, 'section'):
                    section_values = [selected_faculty.section]
                elif hasattr(selected_faculty, 'sections'):
                    # if many-to-many
                    section_values = list(selected_faculty.sections.all().values_list('name', flat=True))
                else:
                    section_values = []

                if section_values:
                    student_names_in_sections = Student.objects.filter(
                        section__in=section_values
                    ).values_list('name', flat=True).distinct()

                    assigned_student_names = StudentSubject.objects.filter(
                        subject_name=selected_subject.name,
                        student_name__in=student_names_in_sections
                    ).values_list('student_name', flat=True).distinct()

                    students_qs = Student.objects.filter(name__in=assigned_student_names).distinct()
                    total_students = students_qs.count()
                else:
                    total_students = 0
                    students_qs = Student.objects.none()

            # --- Feedbacks only for those assigned students (so present_students is consistent) ---
            feedback_list = Feedback.objects.filter(
                faculty_id=faculty_id,
                subject_id=subject_id,
                student__in=students_qs
            )

            present_students = feedback_list.values('student').distinct().count()

            # Questions breakdown (based only on assigned students' feedback)
            questions = FeedbackQuestion.objects.all()
            overall_total = 0

            for idx, question in enumerate(questions, start=1):
                feedbacks = feedback_list.filter(question=question)
                counts = {i: feedbacks.filter(rating=i).count() for i in range(1, 6)}
                total_responses = sum(counts.values())

                weighted_total = (
                    counts[1] * 20 +
                    counts[2] * 40 +
                    counts[3] * 60 +
                    counts[4] * 80 +
                    counts[5] * 100
                )

                avg_score = round(weighted_total / total_responses, 2) if total_responses else 0
                overall_total += avg_score

                feedback_table.append({
                    'sl_no': idx,
                    'question': question.text,
                    'counts': counts,
                    'total': weighted_total,
                    'average': avg_score,
                    'marks': avg_score,
                })

            overall_avg = round(overall_total / len(questions), 2) if questions else 0

            if overall_avg >= 85:
                interpretation = "Excellent"
            elif overall_avg >= 75:
                interpretation = "Very Good"
            elif overall_avg >= 60:
                interpretation = "Good"
            elif overall_avg >= 50:
                interpretation = "Satisfactory"
            else:
                interpretation = "Poor"

    return render(request, 'feedback_app/feedback_report.html', {
        'report': report,
        'selected_faculty': selected_faculty,
        'selected_subject': selected_subject,
        'feedback_table': feedback_table,
        'total_students': total_students,
        'present_students': present_students,
        'overall_avg': overall_avg,
        'interpretation': interpretation,
    })


# Download Excel
import openpyxl
from openpyxl import Workbook
from django.http import HttpResponse
from .models import Student
import xlwt


def download_students_excel(request):
    selected_year = request.GET.get('year')
    students = Student.objects.filter(year=selected_year) if selected_year else Student.objects.all()

    response = HttpResponse(content_type='application/ms-excel')
    response['Content-Disposition'] = f'attachment; filename="students_{selected_year}.xls"'

    wb = xlwt.Workbook(encoding='utf-8')
    ws = wb.add_sheet('Students')

    # Header
    columns = ['USN', 'Name', 'Section', 'User ID', 'Password']
    for col_num, column_title in enumerate(columns):
        ws.write(0, col_num, column_title)

    # Data
    for row_num, student in enumerate(students, start=1):
        ws.write(row_num, 0, student.usn)
        ws.write(row_num, 1, student.name)
        ws.write(row_num, 2, student.section)
        ws.write(row_num, 3, student.generated_userid)
        ws.write(row_num, 4, student.generated_password)

    wb.save(response)
    return response

def download_faculty_excel(request):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Faculty Data"
    ws.append(["Name", "Username"])

    for f in Faculty.objects.all():
        ws.append([f.name, f.user.username])

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename=faculty_data.xlsx'
    wb.save(response)
    return response


from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from .models import Faculty   # ✅ use your Faculty model

def view_faculties(request):
    faculties = Faculty.objects.all().order_by("usn")  # sorted by USN
    return render(request, "feedback_app/view_faculties.html", {"faculties": faculties})

def delete_faculty(request, faculty_id):
    faculty = get_object_or_404(Faculty, id=faculty_id)
    faculty.delete()
    messages.success(request, f"Faculty {faculty.name} deleted successfully!")
    return redirect("view_faculties")

from django.views.decorators.http import require_POST

@require_POST
@login_required
def clear_mappings(request):
    selected_year = request.GET.get("year")
    if selected_year:
        StudentFacultyMap.objects.filter(student__year=selected_year).delete()
    else:
        StudentFacultyMap.objects.all().delete()
    return redirect(f"{request.META.get('HTTP_REFERER', '/dashboard/map-student-faculty/')}")


from io import BytesIO
from django.http import HttpResponse
from xhtml2pdf import pisa
from django.template.loader import get_template
from .models import StudentFacultyMap
import pandas as pd


def download_mappings_pdf(request):
    selected_year = request.GET.get("year")
    mappings = StudentFacultyMap.objects.select_related('student', 'faculty')
    if selected_year:
        mappings = mappings.filter(student__year=selected_year)

    year_label = f"({selected_year})" if selected_year else ""

    html = f"""
    <html>
    <head>
      <style>
        table {{ width: 100%; border-collapse: collapse; font-family: Arial, sans-serif; }}
        th, td {{ border: 1px solid #000; padding: 8px; }}
        th {{ background: #ddd; }}
      </style>
    </head>
    <body>
      <h2>Student to Faculty Mapping {year_label}</h2>
      <table>
        <tr>
          <th>Student Name</th>
          <th>USN</th>
          <th>Faculty</th>
        </tr>
    """

    for m in mappings:
        html += f"""
        <tr>
          <td>{m.student.name}</td>
          <td>{m.student.usn}</td>
          <td>{m.faculty.name}</td>
        </tr>
        """

    html += """
      </table>
    </body>
    </html>
    """

    response = HttpResponse(content_type='application/pdf')
    filename = f'student_faculty_mapping_{selected_year or "all"}.pdf'
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    pisa.CreatePDF(BytesIO(html.encode('UTF-8')), dest=response)
    return response


# ---------------- Excel Export ----------------
def download_mappings_excel(request):
    selected_year = request.GET.get("year")
    mappings = StudentFacultyMap.objects.select_related('student', 'faculty')
    if selected_year:
        mappings = mappings.filter(student__year=selected_year)

    data = [{
        'Student Name': m.student.name,
        'USN': m.student.usn,
        'Faculty': m.faculty.name
    } for m in mappings]

    df = pd.DataFrame(data)

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    filename = f'student_faculty_mapping_{selected_year or "all"}.xlsx'
    response['Content-Disposition'] = f'attachment; filename="{filename}"'

    # Write to Excel
    df.to_excel(response, index=False, sheet_name="Mappings")
    return response
